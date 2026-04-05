from __future__ import annotations

import logging
import re
import subprocess
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from shutil import which
from uuid import uuid4

from app.models.project import ExportArtifact, ExportFormat, ProjectRecord
from app.services.project_service import (
    get_effective_sections,
    get_project,
    get_project_references,
    get_project_title,
    record_export,
)
from core.config import get_settings
from core.exceptions import ExportDependencyError, ExportError
from services.latex_service import generate_latex

logger = logging.getLogger("papereasy.backend.export")
settings = get_settings()

MEDIA_TYPES: dict[ExportFormat, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "latex": "application/x-tex",
}

SECTION_ORDER = (
    ("Abstract", "abstract"),
    ("Introduction", "introduction"),
    ("Methodology", "methodology"),
    ("Conclusion", "conclusion"),
)


@dataclass(frozen=True)
class ExportedFile:
    artifact: ExportArtifact
    file_path: Path
    media_type: str


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value).strip("-").lower()
    return slug or "research-paper"


def _relative_output_path(file_path: Path) -> str:
    return file_path.relative_to(settings.outputs_dir.parent).as_posix()


def _download_url(file_path: Path) -> str:
    return f"/{_relative_output_path(file_path)}"


def _create_artifact(file_path: Path, export_format: ExportFormat) -> ExportArtifact:
    return ExportArtifact(
        id=str(uuid4()),
        format=export_format,
        file_name=file_path.name,
        path=_relative_output_path(file_path),
        download_url=_download_url(file_path),
        created_at=datetime.now(timezone.utc),
    )


def _iter_section_figures(project: ProjectRecord, section: str):
    backend_root = Path(__file__).resolve().parents[1]
    for figure in project.figures:
        if figure.section != section:
            continue

        figure_path = Path(figure.path)
        if not figure_path.is_absolute():
            figure_path = (backend_root / figure.path).resolve()

        if not figure_path.exists():
            logger.warning("Skipping missing figure %s during export", figure.id)
            continue

        yield figure, figure_path


def _write_docx(project: ProjectRecord, output_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ImportError as exc:
        raise ExportDependencyError("python-docx is required for DOCX export.") from exc

    sections = get_effective_sections(project)
    references = get_project_references(project)
    title = get_project_title(project)
    file_name = f"{_slugify(title)}.docx"
    output_path = output_dir / file_name

    document = Document()
    document.add_heading(title, 0)

    for heading, key in SECTION_ORDER:
        document.add_heading(heading, level=1)
        document.add_paragraph(getattr(sections, key))

        for index, (figure, figure_path) in enumerate(_iter_section_figures(project, key), start=1):
            try:
                document.add_picture(str(figure_path), width=Inches(5.75))
            except Exception as exc:
                raise ExportError(
                    f"Unable to embed figure '{figure.original_file_name}' in the DOCX export."
                ) from exc

            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Figure {index}. {figure.caption}")
            run.italic = True

    document.add_heading("References", level=1)
    if references:
        for reference in references:
            document.add_paragraph(reference)
    else:
        document.add_paragraph("No references available.")

    document.save(output_path)
    return output_path


def _run_pdflatex(tex_path: Path) -> Path:
    pdflatex_binary = which(settings.pdflatex_command)
    if pdflatex_binary is None:
        raise ExportDependencyError(
            "pdflatex is not installed or not available on PATH for PDF export."
        )

    command = [
        pdflatex_binary,
        "-interaction=nonstopmode",
        "-halt-on-error",
        tex_path.name,
    ]

    try:
        process = subprocess.run(
            command,
            cwd=tex_path.parent,
            capture_output=True,
            text=True,
            timeout=settings.pdflatex_timeout_seconds,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise ExportError("pdflatex timed out while generating the PDF export.") from exc

    if process.returncode != 0:
        stderr = process.stderr.strip() or process.stdout.strip()
        message = stderr.splitlines()[-1] if stderr else "pdflatex failed to compile the document."
        raise ExportError(f"pdflatex failed: {message}")

    pdf_path = tex_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise ExportError("pdflatex completed without producing a PDF file.")

    for extension in (".aux", ".log", ".out"):
        auxiliary = tex_path.with_suffix(extension)
        if auxiliary.exists():
            auxiliary.unlink(missing_ok=True)

    return pdf_path


def export_project_latex(project_id: str) -> ExportedFile:
    project = get_project(project_id)
    build_result = generate_latex(project, settings.templates_dir, settings.outputs_dir)
    artifact = _create_artifact(build_result.tex_path, "latex")
    record_export(project_id, artifact)
    return ExportedFile(
        artifact=artifact,
        file_path=build_result.tex_path,
        media_type=MEDIA_TYPES["latex"],
    )


def export_project_docx(project_id: str) -> ExportedFile:
    project = get_project(project_id)
    get_effective_sections(project)
    output_dir = settings.outputs_dir / project.id / str(uuid4())
    output_dir.mkdir(parents=True, exist_ok=True)
    file_path = _write_docx(project, output_dir)
    artifact = _create_artifact(file_path, "docx")
    record_export(project_id, artifact)
    logger.info("Generated DOCX export for project %s at %s", project_id, file_path)
    return ExportedFile(artifact=artifact, file_path=file_path, media_type=MEDIA_TYPES["docx"])


def export_project_pdf(project_id: str) -> ExportedFile:
    project = get_project(project_id)
    build_result = generate_latex(project, settings.templates_dir, settings.outputs_dir)
    pdf_path = _run_pdflatex(build_result.tex_path)
    artifact = _create_artifact(pdf_path, "pdf")
    record_export(project_id, artifact)
    logger.info("Generated PDF export for project %s at %s", project_id, pdf_path)
    return ExportedFile(artifact=artifact, file_path=pdf_path, media_type=MEDIA_TYPES["pdf"])
